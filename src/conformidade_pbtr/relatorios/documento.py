"""Relatório em DOCX (formato de tramitação interna)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..modelos import Categoria, Origem, Relatorio, Status

COR = {
    Status.CONFORME: "C6EFCE",
    Status.NAO_CONFORME: "FFC7CE",
    Status.ATENCAO: "FFEB9C",
    Status.VERIFICAR_MANUAL: "DDEBF7",
    Status.NAO_APLICAVEL: "F2F2F2",
}

ROTULO = {
    Status.CONFORME: "Conforme",
    Status.NAO_CONFORME: "Não conforme",
    Status.ATENCAO: "Atenção",
    Status.VERIFICAR_MANUAL: "Verificar",
    Status.NAO_APLICAVEL: "N/A",
}

TITULO_CATEGORIA = {
    Categoria.CHECKLIST: "Conformidade com o roteiro [TI]",
    Categoria.ESTRUTURA: "Estrutura do documento",
    Categoria.NUMERACAO: "Numeração dos itens",
    Categoria.TABELA: "Tabelas e aritmética",
    Categoria.VALOR: "Valores declarados",
    Categoria.ORTOGRAFIA: "Revisão textual",
}


def _sombrear(celula, hex_cor: str) -> None:
    tc_pr = celula._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_cor)
    tc_pr.append(shd)


def _rodape(doc: Document) -> None:
    secao = doc.sections[0]
    p = secao.footer.paragraphs[0]
    p.text = "Relatório gerado pelo MCP Conformidade PB/TR — uso interno"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(7.5)
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def gerar_docx(rel: Relatorio, destino: str | Path) -> str:
    d = Path(destino)
    doc = Document()

    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(10)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2)

    doc.add_heading("Relatório de Análise de Conformidade", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"Projeto Básico / Termo de Referência — {rel.documento.nome}").bold = True

    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Light List Accent 1"
    for rotulo, valor in [
        ("Documento", rel.documento.nome),
        ("Tipo", rel.documento.tipo),
        ("Páginas", str(rel.documento.paginas)),
        ("Tabelas detectadas", str(len(rel.documento.tabelas))),
        ("Contexto inferido", ", ".join(rel.documento.tags_contexto)),
        ("Checklist", f"v{rel.versao_checklist} — Roteiro [TI] PB/TR"),
        ("Gerado em", rel.gerado_em),
    ]:
        linha = meta.add_row().cells
        linha[0].text = rotulo
        linha[1].text = valor
        linha[0].paragraphs[0].runs[0].bold = True

    # --------------------------------------------------------- sumário
    doc.add_heading("1. Sumário executivo", level=1)
    r = rel.resumo
    p = doc.add_paragraph()
    run = p.add_run(f"Índice de conformidade: {r.indice_conformidade}%")
    run.bold = True
    run.font.size = Pt(14)

    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    hdr[0].text = "Situação"
    hdr[1].text = "Quantidade"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
        _sombrear(c, "D9D9D9")
    for status, qtd in [
        (Status.CONFORME, r.conforme),
        (Status.NAO_CONFORME, r.nao_conforme),
        (Status.ATENCAO, r.atencao),
        (Status.VERIFICAR_MANUAL, r.verificar_manual),
        (Status.NAO_APLICAVEL, r.nao_aplicavel),
    ]:
        linha = t.add_row().cells
        linha[0].text = ROTULO[status]
        linha[1].text = str(qtd)
        _sombrear(linha[0], COR[status])

    doc.add_paragraph()
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    h = t2.rows[0].cells
    h[0].text = "Verificação automática"
    h[1].text = "Ocorrências"
    for c in h:
        c.paragraphs[0].runs[0].bold = True
        _sombrear(c, "D9D9D9")
    for rotulo, qtd in [
        ("Erros de numeração", r.erros_numeracao),
        ("Divergências em tabelas/valores", r.erros_tabela),
        ("Apontamentos de revisão textual (regra)", r.erros_ortografia),
        ("Sugestões da revisão pelo agente", r.sugestoes_ia),
    ]:
        linha = t2.add_row().cells
        linha[0].text = rotulo
        linha[1].text = str(qtd)

    if rel.avisos:
        doc.add_heading("Avisos da análise", level=2)
        for av in rel.avisos:
            doc.add_paragraph(av, style="List Bullet")

    # ------------------------------------------------------ pendências
    criticas = [
        a for a in rel.pendencias()
        if a.severidade.value in ("critico", "alto") and a.status == Status.NAO_CONFORME
    ]
    if criticas:
        doc.add_heading("2. Pendências prioritárias", level=1)
        tp = doc.add_table(rows=1, cols=4)
        tp.style = "Table Grid"
        cab = ["ID", "Seção", "Pendência", "Recomendação"]
        for i, c in enumerate(tp.rows[0].cells):
            c.text = cab[i]
            c.paragraphs[0].runs[0].bold = True
            _sombrear(c, "D9D9D9")
        for a in criticas[:30]:
            linha = tp.add_row().cells
            linha[0].text = a.id
            linha[1].text = a.secao
            linha[2].text = a.titulo
            linha[3].text = a.orientacao or "—"

    # -------------------------------------------------------- detalhes
    n_secao = 3
    for cat in [
        Categoria.CHECKLIST,
        Categoria.ESTRUTURA,
        Categoria.NUMERACAO,
        Categoria.TABELA,
        Categoria.VALOR,
        Categoria.ORTOGRAFIA,
    ]:
        itens = rel.por_categoria(cat)
        if not itens:
            continue
        doc.add_heading(f"{n_secao}. {TITULO_CATEGORIA[cat]}", level=1)
        n_secao += 1

        if cat == Categoria.CHECKLIST:
            tab = doc.add_table(rows=1, cols=5)
            tab.style = "Table Grid"
            cab = ["ID", "Item do roteiro", "Situação", "Constatação", "Recomendação"]
            larguras = [Cm(1.8), Cm(5.5), Cm(2.2), Cm(4.5), Cm(4.5)]
            for i, c in enumerate(tab.rows[0].cells):
                c.text = cab[i]
                c.width = larguras[i]
                c.paragraphs[0].runs[0].bold = True
                _sombrear(c, "D9D9D9")
            for a in sorted(itens, key=lambda x: x.id):
                linha = tab.add_row().cells
                valores = [
                    a.id,
                    f"{a.secao}\n{a.titulo}",
                    ROTULO[a.status],
                    a.encontrado or "—",
                    "—"
                    if a.status in (Status.CONFORME, Status.NAO_APLICAVEL)
                    else (a.orientacao or "—"),
                ]
                for i, v in enumerate(valores):
                    linha[i].text = v
                    linha[i].width = larguras[i]
                    for par in linha[i].paragraphs:
                        for run in par.runs:
                            run.font.size = Pt(8)
                _sombrear(linha[2], COR[a.status])
        elif cat == Categoria.ORTOGRAFIA:
            for origem, subtitulo, nota in (
                (
                    Origem.DETERMINISTICO,
                    "Verificações determinísticas",
                    "Regras exatas e reprodutíveis: o mesmo documento produz sempre o "
                    "mesmo resultado.",
                ),
                (
                    Origem.IA,
                    "Sugestões da revisão pelo agente",
                    "Leitura do texto por modelo de linguagem. Cada trecho citado foi "
                    "conferido contra o documento, mas a revisão não é reprodutível — "
                    "trate como sugestão, não como constatação.",
                ),
            ):
                subitens = rel.por_categoria(cat, origem)
                if not subitens:
                    continue
                doc.add_heading(subtitulo, level=2)
                aviso = doc.add_paragraph()
                aviso.add_run(nota).italic = True
                for a in subitens:
                    par = doc.add_paragraph()
                    par.add_run(f"[{a.id}] {a.titulo}").bold = True
                    for rotulo, valor in [
                        ("Trecho", a.encontrado),
                        ("Sugestão", a.esperado),
                        ("Localização", ", ".join(x for x in ((f"item {a.item}" if a.item else ""), (f"p. {a.pagina}" if a.pagina else "")) if x)),
                        ("Observação", a.descricao),
                        ("Recomendação", a.orientacao),
                    ]:
                        if valor:
                            item = doc.add_paragraph(style="List Bullet")
                            item.add_run(f"{rotulo}: ").bold = True
                            item.add_run(valor)
        else:
            for a in itens:
                if a.status == Status.CONFORME and cat != Categoria.TABELA:
                    continue
                p = doc.add_paragraph()
                p.add_run(f"[{a.id}] {a.titulo}").bold = True
                p.add_run(f"  ({ROTULO[a.status]} · {a.severidade.value})").italic = True
                if a.descricao:
                    doc.add_paragraph(a.descricao)
                for rotulo, valor in [
                    ("Esperado", a.esperado),
                    ("Encontrado", a.encontrado),
                    ("Trecho", (f"(item {a.item}, p. {a.pagina}) " if a.item else f"(p. {a.pagina or '—'}) ") + a.evidencia[:300] if a.evidencia else ""),
                    ("Recomendação", a.orientacao),
                ]:
                    if valor:
                        item = doc.add_paragraph(style="List Bullet")
                        item.add_run(f"{rotulo}: ").bold = True
                        item.add_run(valor)

    doc.add_paragraph()
    nota = doc.add_paragraph()
    nota.add_run(
        "As verificações automáticas não substituem a análise do parecerista. "
        "Itens marcados como “Verificar” exigem conferência humana, notadamente "
        "a aderência entre a Aba Itens e as quantidades do PB."
    ).italic = True

    _rodape(doc)
    doc.save(str(d))
    return str(d)
